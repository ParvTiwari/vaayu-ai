"""
Render docs/DETAILED_SUBMISSION_DOCUMENT.md to dist/Vaayu-AI-Detailed-Submission-Document.docx.

A minimal, dependency-light markdown -> docx converter tailored to this one
document's structure (headers, bold, tables, bullet/numbered lists, code
blocks, horizontal rules, links). Not a general-purpose markdown renderer —
if the source doc starts using markdown features not listed above, extend
this rather than reaching for a heavier dependency.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "DETAILED_SUBMISSION_DOCUMENT.md"
OUT = ROOT / "dist" / "Vaayu-AI-Detailed-Submission-Document.docx"

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_CODE_RE = re.compile(r"`([^`]+)`")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Add `text` to `paragraph`, honoring **bold**, `code`, and [text](url) spans."""
    # Normalize links to "text (url)" before the bold/code pass, so URLs survive.
    text = _LINK_RE.sub(lambda m: f"{m.group(1)} ({m.group(2)})", text)

    # Split on bold and code spans, tagging each token with its style.
    tokens: list[tuple[str, str]] = []
    pos = 0
    combined = re.compile(r"(\*\*(?:.+?)\*\*|`(?:[^`]+)`)")
    for m in combined.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos : m.start()]))
        chunk = m.group(0)
        if chunk.startswith("**"):
            tokens.append(("bold", chunk[2:-2]))
        else:
            tokens.append(("code", chunk[1:-1]))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))

    for kind, content in tokens:
        run = paragraph.add_run(content)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c >= n_cols:
                continue
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text.strip())
            if r == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def _parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def convert(md_text: str) -> Document:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = md_text.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    table_buffer: list[str] = []

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            data_lines = [ln for ln in table_buffer if not re.match(r"^\s*\|?\s*:?-{2,}", ln)]
            rows = [_parse_table_row(ln) for ln in data_lines]
            _add_table(doc, rows)
            table_buffer = []

    while i < len(lines):
        line = lines[i]

        # Fenced code blocks
        if line.strip().startswith("```"):
            if not in_code_block:
                flush_table()
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(8)
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Markdown table rows
        if stripped.startswith("|"):
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            border_run = p.add_run("_" * 60)
            border_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            i += 1
            continue

        # Headers
        m = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=level if level <= 3 else 3)
            _add_inline_runs(heading, text)
            i += 1
            continue

        # Blockquote (used for a callout in this doc)
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _add_inline_runs(p, stripped.lstrip("> ").strip())
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m.group(2))
            i += 1
            continue

        # Bullet list (including nested "  - ")
        m = re.match(r"^-\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(1))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    flush_table()
    return doc


def main() -> None:
    if not SRC.exists():
        raise SystemExit(f"{SRC} not found")
    md_text = SRC.read_text(encoding="utf-8")
    doc = convert(md_text)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
